from faker import Faker
from docx import Document
import random
import os
import json

fake = Faker()

class SyntheticResumeGenerator:
    def __init__(self, taxonomy_path: str):
        with open(taxonomy_path, 'r') as f:
            self.taxonomy = json.load(f)
        self.categories = list(self.taxonomy.keys())

    def generate_resume(self, output_path: str):
        doc = Document()
        
        # Name and Header
        name = fake.name()
        doc.add_heading(name, 0)
        doc.add_paragraph(f"{fake.email()} | {fake.phone_number()} | {fake.city()}")

        # Summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(fake.paragraph(nb_sentences=3))

        # Skills
        doc.add_heading('Skills', level=1)
        # Select random skills from taxonomy
        selected_skills = []
        for _ in range(3):
            cat = random.choice(self.categories)
            selected_skills.extend(random.sample(self.taxonomy[cat], k=min(3, len(self.taxonomy[cat]))))
        
        doc.add_paragraph(", ".join(list(set(selected_skills))))

        # Experience
        doc.add_heading('Experience', level=1)
        for _ in range(2):
            doc.add_heading(fake.company(), level=2)
            doc.add_paragraph(f"{fake.job()} | {fake.date()}")
            doc.add_paragraph(fake.paragraph(nb_sentences=4))

        # Education
        doc.add_heading('Education', level=1)
        doc.add_paragraph(f"{fake.company()} | Degree in {random.choice(self.categories)}")

        doc.save(output_path)
        return name, selected_skills

if __name__ == "__main__":
    generator = SyntheticResumeGenerator("resume-intelligence/data/skills_taxonomy.json")
    os.makedirs("resume-intelligence/data/raw", exist_ok=True)
    
    for i in range(5):
        name, skills = generator.generate_resume(f"resume-intelligence/data/raw/resume_{i}.docx")
        print(f"Generated resume for {name} with skills: {skills}")
